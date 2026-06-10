"""
Document service for handling document uploads, text extraction, and the full
knowledge-base pipeline (chunk → embed → store in ChromaDB).
"""

import io
from typing import Optional, Tuple
from uuid import UUID
from pathlib import Path
from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession
from loguru import logger

from app.repositories.document import DocumentRepository
from app.repositories.document_chunk import DocumentChunkRepository
from app.models.document import Document, DocumentType
from app.core.config import settings


class DocumentService:
    """Service for document handling and text extraction."""

    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def __init__(self, db: AsyncSession):
        self.db = db
        self.repo = DocumentRepository(db)
        self.chunk_repo = DocumentChunkRepository(db)
        self.upload_dir = Path(settings.UPLOAD_DIRECTORY)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    async def upload_document(
        self,
        user_id: UUID,
        file: UploadFile,
    ) -> Tuple[Document, str]:
        """
        Validate, save, and extract text from an uploaded file.
        Returns (Document, extracted_text). Caller schedules pipeline.
        """
        try:
            self._validate_file(file)

            content = await file.read()
            file_size = len(content)

            if file_size > self.MAX_FILE_SIZE:
                raise ValueError(
                    f"File size {file_size / 1024 / 1024:.1f} MB exceeds "
                    f"maximum of {self.MAX_FILE_SIZE / 1024 / 1024:.0f} MB"
                )

            file_ext = Path(file.filename).suffix.lower()
            document_type = self._get_document_type(file_ext)

            file_path = await self._save_file(user_id, file.filename, content)

            # Extract FULL text (not just a preview)
            full_text = await self._extract_full_text(document_type, content)

            # Strip null bytes and other non-PostgreSQL-safe characters
            if full_text:
                full_text = full_text.replace("\x00", "").strip()
                if not full_text:
                    full_text = None

            document = await self.repo.create_document(
                user_id=user_id,
                filename=file.filename,
                document_type=document_type,
                file_path=str(file_path),
                file_size=file_size,
                content_preview=full_text[:500] if full_text else None,
                extracted_text=full_text,
            )

            logger.info(
                f"Document uploaded: {file.filename} (ID: {document.id}, "
                f"text_len={len(full_text) if full_text else 0})"
            )
            return document, full_text or ""

        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            raise

    async def ingest_url(self, user_id: UUID, url: str, title: Optional[str] = None) -> Document:
        """
        Scrape a webpage, extract its text, and store it as a knowledge-base document.
        Returns the Document so the caller can schedule the embedding pipeline.
        """
        import requests as _req
        from bs4 import BeautifulSoup

        logger.info(f"Scraping URL: {url}")
        try:
            resp = _req.get(url, timeout=30, headers={"User-Agent": "UniConnect/1.0"})
            resp.raise_for_status()
        except Exception as e:
            raise ValueError(f"Could not fetch URL: {e}")

        soup = BeautifulSoup(resp.text, "lxml")

        # Remove boilerplate tags
        for tag in soup(["script", "style", "nav", "footer", "header",
                         "aside", "noscript", "iframe", "form"]):
            tag.decompose()

        page_title = title or (soup.title.string.strip() if soup.title else url)
        safe_title = "".join(c if c.isalnum() or c in " .-_" else "_" for c in page_title)[:120]

        # Extract clean text
        text = soup.get_text(separator="\n", strip=True)
        # Collapse excessive blank lines
        import re as _re
        text = _re.sub(r"\n{3,}", "\n\n", text).strip()

        if not text or len(text) < 50:
            raise ValueError("Page contained no extractable text.")

        logger.info(f"Scraped {len(text)} chars from: {url}")

        # Save as a virtual .txt file so the pipeline works unchanged
        filename = f"{safe_title}.txt"
        file_content = f"Source: {url}\nTitle: {page_title}\n\n{text}".encode("utf-8")
        file_path = await self._save_file(user_id, filename, file_content)

        document = await self.repo.create_document(
            user_id=user_id,
            filename=filename,
            document_type=DocumentType.URL,
            file_path=str(file_path),
            file_size=len(file_content),
            content_preview=text[:500],
            extracted_text=text,
            source_url=url,
        )
        logger.info(f"URL ingested: {url} → document {document.id}")
        return document

    async def get_document(self, document_id: UUID) -> Optional[Document]:
        return await self.repo.get_document_by_id(document_id)

    async def get_user_documents(self, user_id: UUID) -> list:
        return await self.repo.get_documents_by_user(user_id)

    async def get_all_documents(self) -> list:
        return await self.repo.get_all_documents()

    async def delete_document(self, document_id: UUID, user_id: UUID) -> bool:
        try:
            document = await self.repo.get_document_by_id(document_id)
            if not document:
                return False
            if document.user_id != user_id:
                logger.warning(f"User {user_id} not authorized to delete {document_id}")
                return False

            try:
                file_path = Path(document.file_path)
                if file_path.exists():
                    file_path.unlink()
            except Exception as e:
                logger.warning(f"Could not delete file on disk: {e}")

            return await self.repo.delete_document(document_id)

        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            return False

    # ------------------------------------------------------------------
    # Validation helpers
    # ------------------------------------------------------------------

    def _validate_file(self, file: UploadFile) -> None:
        if not file.filename:
            raise ValueError("Filename is required")
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in self.ALLOWED_EXTENSIONS:
            raise ValueError(
                f"File type '{file_ext}' not allowed. "
                f"Allowed: {', '.join(sorted(self.ALLOWED_EXTENSIONS))}"
            )
        if file.size and file.size > self.MAX_FILE_SIZE:
            raise ValueError(
                f"File size exceeds maximum of {self.MAX_FILE_SIZE / 1024 / 1024:.0f} MB"
            )

    @staticmethod
    def _get_document_type(file_ext: str) -> DocumentType:
        return {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".txt": DocumentType.TXT,
        }.get(file_ext.lower(), DocumentType.TXT)

    # ------------------------------------------------------------------
    # Storage helpers
    # ------------------------------------------------------------------

    async def _save_file(self, user_id: UUID, filename: str, content: bytes) -> Path:
        user_dir = self.upload_dir / str(user_id)
        user_dir.mkdir(parents=True, exist_ok=True)
        safe_filename = f"{user_id}_{filename}"
        file_path = user_dir / safe_filename
        with open(file_path, "wb") as f:
            f.write(content)
        logger.info(f"File saved: {file_path}")
        return file_path

    # ------------------------------------------------------------------
    # Full-text extraction (NO page/char limits — complete document)
    # ------------------------------------------------------------------

    async def _extract_full_text(
        self, document_type: DocumentType, content: bytes
    ) -> Optional[str]:
        try:
            if document_type == DocumentType.TXT:
                return content.decode("utf-8", errors="ignore").strip()

            if document_type == DocumentType.PDF:
                return await self._extract_pdf_text(content)

            if document_type == DocumentType.DOCX:
                return await self._extract_docx_text(content)

        except Exception as e:
            logger.warning(f"Text extraction failed: {e}")
        return None

    async def _extract_pdf_text(self, content: bytes) -> Optional[str]:
        """
        Three-tier PDF text extraction for maximum reliability:
          1. PyPDF2        — instant, works on all text-based PDFs
          2. Tesseract OCR — local, no API limits, handles any scanned PDF
          3. Gemini Vision — cloud fallback if Tesseract produces poor results
        """
        # ── Tier 1: PyPDF2 (text-based PDFs) ────────────────────────────────
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            parts = [page.extract_text() or "" for page in reader.pages]
            result = "\n".join(parts).strip()
            avg_chars = len(result) / max(len(reader.pages), 1)
            if avg_chars >= 80:
                logger.info(f"PDF: text-based ({len(result)} chars, {avg_chars:.0f}/page) — PyPDF2")
                return result
            logger.info(f"PDF: sparse ({avg_chars:.0f} chars/page) — switching to OCR")
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")

        # ── Convert pages to images (shared by both OCR tiers) ───────────────
        try:
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(content, dpi=200)
            logger.info(f"Converted PDF to {len(images)} page image(s) at 200 DPI")
        except Exception as e:
            logger.error(f"PDF→image conversion failed: {e}")
            return None

        # ── Tier 2: Tesseract OCR (local, no rate limits) ────────────────────
        tesseract_text = await self._ocr_with_tesseract(images)
        if tesseract_text and len(tesseract_text) >= 50:
            logger.info(f"Tesseract OCR complete: {len(tesseract_text)} chars extracted")
            return tesseract_text
        logger.warning("Tesseract produced little/no text — trying Gemini Vision as fallback")

        # ── Tier 3: Gemini Vision (cloud, handles complex layouts) ───────────
        return await self._extract_pdf_with_vision(images)

    async def _ocr_with_tesseract(self, images: list) -> Optional[str]:
        """Run Tesseract OCR on a list of PIL images. Runs in thread pool to avoid blocking."""
        import asyncio as _asyncio
        import concurrent.futures

        def _run_tesseract(images):
            try:
                import pytesseract
                parts = []
                for i, image in enumerate(images, start=1):
                    # Enhance image for better OCR accuracy
                    try:
                        from PIL import ImageFilter, ImageEnhance
                        image = image.convert("L")                          # grayscale
                        image = ImageEnhance.Contrast(image).enhance(2.0)  # boost contrast
                        image = image.filter(ImageFilter.SHARPEN)           # sharpen edges
                    except Exception:
                        pass
                    text = pytesseract.image_to_string(
                        image,
                        config="--psm 1 --oem 3",  # auto page layout + LSTM engine
                    ).replace("\x00", "").strip()
                    if text:
                        parts.append(f"[Page {i}]\n{text}")
                    logger.info(f"Tesseract page {i}: {len(text)} chars")
                return "\n\n".join(parts).strip()
            except Exception as e:
                logger.warning(f"Tesseract OCR error: {e}")
                return None

        loop = _asyncio.get_event_loop()
        with concurrent.futures.ThreadPoolExecutor() as pool:
            return await loop.run_in_executor(pool, _run_tesseract, images)

    async def _extract_pdf_with_vision(self, images: list) -> Optional[str]:
        """
        Gemini Vision OCR — sends each page image to Gemini.
        Used as fallback when Tesseract yields poor results.
        Retries on 429 with exponential backoff, tries multiple models.
        """
        import asyncio as _asyncio
        import base64
        import requests as _requests

        _VISION_MODELS = ["gemini-2.0-flash", "gemini-2.0-flash-lite", "gemini-2.5-flash"]

        try:
            all_text: list[str] = []
            for page_num, image in enumerate(images, start=1):
                img_buffer = io.BytesIO()
                image.save(img_buffer, format="JPEG", quality=85)
                img_b64 = base64.b64encode(img_buffer.getvalue()).decode("utf-8")

                payload = {
                    "contents": [{"parts": [
                        {"inline_data": {"mime_type": "image/jpeg", "data": img_b64}},
                        {"text": (
                            "Extract ALL text from this document image exactly as written. "
                            "Preserve the original structure — headings, paragraphs, lists, tables. "
                            "Output only the extracted text, nothing else."
                        )},
                    ]}],
                    "generationConfig": {"temperature": 0.0, "maxOutputTokens": 4096},
                }

                page_text = ""
                for model in _VISION_MODELS:
                    success = False
                    for attempt in range(3):
                        try:
                            resp = _requests.post(
                                f"https://generativelanguage.googleapis.com/v1beta/models/"
                                f"{model}:generateContent",
                                params={"key": settings.GEMINI_API_KEY},
                                json=payload,
                                timeout=90,
                            )
                            if resp.status_code == 429:
                                wait = 20 * (attempt + 1)
                                logger.warning(f"Vision OCR page {page_num} ({model}): 429 — waiting {wait}s")
                                await _asyncio.sleep(wait)
                                continue
                            resp.raise_for_status()
                            page_text = resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
                            success = True
                            break
                        except Exception as e:
                            if attempt < 2:
                                await _asyncio.sleep(10)
                            else:
                                logger.warning(f"Vision OCR page {page_num} ({model}) failed: {e}")
                    if success:
                        break

                if page_text:
                    all_text.append(f"[Page {page_num}]\n{page_text}")
                    logger.info(f"Vision OCR: page {page_num} done ({len(page_text)} chars)")

                if page_num < len(images):
                    await _asyncio.sleep(3)

            result = "\n\n".join(all_text).strip()
            if result:
                logger.info(f"Gemini Vision OCR complete: {len(result)} chars")
                return result
            logger.warning("Gemini Vision OCR returned no text")
            return None

        except Exception as e:
            logger.error(f"Gemini Vision OCR failed: {e}")
            return None

    async def _extract_docx_text(self, content: bytes) -> Optional[str]:
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(content))
            parts = []
            for paragraph in doc.paragraphs:   # ALL paragraphs, no limit
                if paragraph.text.strip():
                    parts.append(paragraph.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            result = "\n".join(parts).strip()
            return result or None
        except Exception as e:
            logger.warning(f"DOCX extraction error: {e}")
            return None


# ---------------------------------------------------------------------------
# Standalone pipeline — runs as a FastAPI BackgroundTask with its own session
# ---------------------------------------------------------------------------

async def process_document_pipeline(document_id: UUID) -> None:
    """
    Full knowledge-base pipeline for a document:
      1. Chunk the extracted text
      2. Delete any previously stored chunks (safe re-processing)
      3. Persist new chunks to PostgreSQL
      4. Generate Gemini embeddings
      5. Upsert vectors into ChromaDB
      6. Mark document as completed (or failed)

    Creates its own DB session so it can run safely as a background task
    after the HTTP request that triggered it has already returned.
    """
    from app.db.database import AsyncSessionLocal
    from app.services.chunking import split_text
    from app.services.embeddings import EmbeddingService
    from app.services.chromadb_client import ChromaClient

    async with AsyncSessionLocal() as db:
        doc_repo = DocumentRepository(db)
        chunk_repo = DocumentChunkRepository(db)

        try:
            doc = await doc_repo.get_document_by_id(document_id)
            if not doc:
                logger.error(f"Pipeline: document {document_id} not found")
                return

            if not doc.extracted_text:
                logger.error(f"Pipeline: no extracted text for document {document_id}")
                await doc_repo.update_document_status(document_id, "failed")
                return

            text_len = len(doc.extracted_text)
            logger.info(f"Pipeline: starting for {doc.filename} ({text_len} chars)")
            await doc_repo.update_document_status(document_id, "processing")

            # 1. Split into chunks — treat very short text as a single chunk
            if text_len < settings.CHUNK_SIZE:
                chunks_text = [doc.extracted_text]
                logger.info(f"Pipeline: text shorter than chunk size, using as single chunk")
            else:
                chunks_text = split_text(
                    doc.extracted_text,
                    settings.CHUNK_SIZE,
                    settings.CHUNK_OVERLAP,
                )

            if not chunks_text:
                logger.error(f"Pipeline: chunking produced no chunks for {document_id}")
                await doc_repo.update_document_status(document_id, "failed")
                return

            logger.info(f"Pipeline: {len(chunks_text)} chunks created")

            # 2. Remove old chunks (idempotent re-processing)
            await chunk_repo.delete_by_document(document_id)

            # 3. Persist chunks
            db_chunks = await chunk_repo.create_chunks(document_id, chunks_text)

            # 4 & 5. Embed + store in ChromaDB
            logger.info(f"Pipeline: starting embedding for {len(db_chunks)} chunks")
            embed_service = EmbeddingService()
            chroma = ChromaClient()

            texts = [c.content for c in db_chunks]
            ids = [str(c.id) for c in db_chunks]
            metadatas = [
                {
                    "source_type": "document",
                    "document_id": str(document_id),
                    "filename": doc.filename,
                    "chunk_index": c.chunk_index,
                }
                for c in db_chunks
            ]

            logger.info(f"Pipeline: calling Gemini embedding API...")
            vectors = embed_service.embed_texts(texts)
            logger.info(f"Pipeline: embeddings done, upserting to ChromaDB...")
            chroma.upsert(
                settings.CHROMA_COLLECTION_NAME,
                ids=ids,
                embeddings=vectors,
                metadatas=metadatas,
                documents=texts,
            )

            # 6. Mark complete
            await doc_repo.update_document_status(
                document_id, "completed", total_chunks=len(db_chunks)
            )
            logger.info(
                f"Pipeline complete: document {document_id}, "
                f"{len(db_chunks)} chunks indexed"
            )

        except Exception as e:
            logger.error(f"Pipeline failed for document {document_id}: {e}")
            try:
                await doc_repo.update_document_status(document_id, "failed")
            except Exception:
                pass
